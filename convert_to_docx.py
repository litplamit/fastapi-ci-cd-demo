"""
Convert the FastAPI guide markdown to DOCX (Word document)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

# Read the markdown file
md_file = Path(r"C:\Users\Amit_Jirange\.gemini\antigravity\brain\0b83aa23-19fe-4e56-baaa-74dc544296b7\fastapi_complete_guide.md")
output_docx = Path(r"d:\Amit\MyFASTAPI\FastAPI_Complete_Guide.docx")

with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 150, 136)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(0, 121, 107)
    return heading

def add_code_block(code_text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Add background color effect with shading
    return para

def add_blockquote(text):
    """Add a blockquote"""
    para = doc.add_paragraph(text)
    para.paragraph_format.left_indent = Inches(0.5)
    para.runs[0].font.italic = True
    para.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    return para

# Process markdown content
lines = content.split('\n')
in_code_block = False
code_buffer = []
in_blockquote = False

for line in lines:
    # Handle code blocks
    if line.startswith('```'):
        if in_code_block:
            # End of code block
            add_code_block('\n'.join(code_buffer))
            code_buffer = []
            in_code_block = False
        else:
            # Start of code block
            in_code_block = True
        continue
    
    if in_code_block:
        code_buffer.append(line)
        continue
    
    # Handle headings
    if line.startswith('# '):
        add_heading(line[2:], level=1)
    elif line.startswith('## '):
        add_heading(line[3:], level=2)
    elif line.startswith('### '):
        add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        add_heading(line[5:], level=4)
    
    # Handle blockquotes
    elif line.startswith('> '):
        add_blockquote(line[2:])
    
    # Handle horizontal rules
    elif line.strip() == '---':
        doc.add_paragraph('_' * 80)
    
    # Handle lists
    elif line.startswith('- ') or line.startswith('* '):
        para = doc.add_paragraph(line[2:], style='List Bullet')
    elif re.match(r'^\d+\. ', line):
        para = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
    
    # Handle regular paragraphs
    elif line.strip():
        # Skip mermaid diagrams and special markdown
        if 'mermaid' not in line and not line.startswith('[') and not line.startswith('!'):
            para = doc.add_paragraph(line)

# Add title page
doc.paragraphs[0].insert_paragraph_before('Complete FastAPI Application Guide', style='Title')
doc.paragraphs[1].insert_paragraph_before('From Zero to Production - A Beginner-Friendly Deep Dive', style='Subtitle')

# Save document
doc.save(output_docx)

print(f"✓ Word document created: {output_docx}")
print(f"\nYou can now:")
print(f"1. Open {output_docx.name} in Microsoft Word")
print(f"2. Save as PDF from Word (File > Save As > PDF)")
